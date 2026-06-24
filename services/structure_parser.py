"""文档结构解析器。

多层降级链（零 LLM 调用）：
1. docx → python-docx 原生 Heading 样式
2. MinerU Markdown → 正则解析 # / ## / ###
3. 纯文本 → 正则解析编号标题
4. 都没 → 整篇作为单章节
"""

import re
from typing import Optional

from models.audit_document import DocumentStructure, Chapter, Clause


# ── 主入口 ──────────────────────────────────────────────────────────────────

def extract_structure(parsed_content: str, file_type: str = "", file_path: str = "") -> DocumentStructure:
    """多层降级链提取文档结构，零 LLM 调用。

    策略：
    1. 正则解析 Markdown 标题（# / ## / ###）— 覆盖所有格式，不受 Word 样式限制
    2. docx 原生 Heading 样式—补充条款层级（如果正则找不到 H3）
    3. 整篇作为单章节—兜底

    Returns:
        始终返回有效结构，最后一层兜底为单章节。
    """

    # ── 第 1 层：正则解析（不受 Word 样式限制，识别所有 # 标题）──
    structure = _parse_by_regex(parsed_content)
    if structure and structure.total_clauses > 0:
        return structure
    if structure and len(structure.chapters) > 0:
        return structure

    # ── 第 2 层：docx 原生 Heading 样式（补充，仅当正则完全失败时）──
    if file_type == "docx" and file_path:
        try:
            structure = _extract_from_docx_styles(file_path, parsed_content)
            if structure and len(structure.chapters) > 1:
                return structure
        except Exception:
            pass

    # ── 第 3 层：整篇作为单章节 ──
    return DocumentStructure(
        title="全文",
        chapters=[Chapter(title="全文", text=parsed_content)],
        total_clauses=0,
    )


# ── 第 1 层：docx Heading 样式 ──────────────────────────────────────────────

def _extract_from_docx_styles(file_path: str, parsed_content: str) -> Optional[DocumentStructure]:
    """从 docx 原生 Heading 样式提取章节结构，同时将章节文本映射回 parsed_content。"""
    from docx import Document as DocxDocument

    docx = DocxDocument(file_path)
    chapters: list[Chapter] = []
    current_chapter: Optional[Chapter] = None
    clause_counter = 0

    for para in docx.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "Normal"

        # Heading 1 → 新章节
        if "Heading 1" in style_name:
            current_chapter = Chapter(title=text)
            chapters.append(current_chapter)
        # Heading 2 / Heading 3 → 新条款
        elif "Heading" in style_name and current_chapter is not None:
            clause_counter += 1
            num = _extract_number(text)
            current_chapter.clauses.append(Clause(
                number=num or f"{clause_counter}",
                text=text[:200],
            ))

    # 如果 docx 没用原生 Heading 样式（全 Normal），返回 None 走降级
    if len(chapters) <= 1:
        for para in docx.paragraphs:
            if para.style and "Heading" in para.style.name:
                return _map_chapter_texts(chapters, parsed_content)
        return None  # 没有 Heading 样式 → 降级

    return _map_chapter_texts(chapters, parsed_content)


def _map_chapter_texts(chapters: list[Chapter], parsed_content: str) -> DocumentStructure:
    """将章节标题映射回 parsed_content 中的文本位置，填充 chapter.text。"""
    total_clauses = 0
    for i, ch in enumerate(chapters):
        total_clauses += len(ch.clauses)
        # 找标题在 parsed_content 中的位置，提取到下一个标题为止
        start = parsed_content.find(ch.title)
        if start >= 0:
            # 找下一个章节标题的位置
            end = len(parsed_content)
            for j in range(i + 1, len(chapters)):
                next_start = parsed_content.find(chapters[j].title, start + 1)
                if next_start > start:
                    end = next_start
                    break
            ch.text = parsed_content[start:end].strip()
        else:
            # 标题在 parsed_content 中找不到（格式不同），用全文
            ch.text = parsed_content[:5000]

    total = sum(len(ch.clauses) for ch in chapters)
    return DocumentStructure(
        title=chapters[0].title if chapters else "",
        chapters=chapters,
        total_clauses=total,
    )


# ── 第 2 层：正则解析 ───────────────────────────────────────────────────────

def _parse_by_regex(content: str) -> Optional[DocumentStructure]:
    """用正则从 Markdown / 纯文本中提取章节和条款，同时保留章节原文。"""
    lines = content.split("\n")
    chapters: list[Chapter] = []
    current_chapter: Optional[Chapter] = None
    current_clause_num: Optional[str] = None
    current_clause_text: list[str] = []
    chapter_start_line = 0  # 当前章节在 lines 中的起始行

    # 扫描文档，判断最高标题级别
    # 标准规范文件通常用 ##（H2）作为章节标题（如 "## 1 总则"），
    # 而 MinerU 输出的文档用 #（H1）作为章节标题。
    # 这里动态适应，确保两种格式都能正确识别章节边界。
    cleaned_lines = _flatten_html_tables(lines)
    has_h1 = any(re.match(r'^#\s+\S', line) for line in cleaned_lines)
    chapter_heading_level = 1 if has_h1 else 2
    chapter_re = re.compile(r'^#{' + str(chapter_heading_level) + r'}\s+(.+)$')

    # 中文文档标题模式（无 Markdown 标记时，常见于 DOCX 提取的纯文本）
    # 章节级：第X部分 / 第X章 / 一、 / （一）
    # 子标题级：1. 概述 / 1.1 项目概况
    cn_part_re = re.compile(r'^第[一二三四五六七八九十百]+[部分章节编篇]\s*(.*)')
    cn_num_re = re.compile(r'^[一二三四五六七八九十]+[、.．]\s*(.*)')
    cn_paren_re = re.compile(r'^（[一二三四五六七八九十]+）\s*(.*)')
    # 用于判断是否应该走中文标题路径（文档中是否有 Markdown 标题）
    has_markdown_headings = has_h1 or any(re.match(r'^#{2,3}\s+\S', line) for line in cleaned_lines)

    for idx, line in enumerate(cleaned_lines):
        stripped = line.strip()
        if not stripped:
            continue

        # 章节边界（H1 或 H2，取决于文档的标题层级）
        chapter_match = chapter_re.match(stripped)
        if chapter_match:
            _finalize_clause(current_chapter, current_clause_num, current_clause_text)
            # 填充上一章节的 text
            if current_chapter is not None:
                current_chapter.text = "\n".join(lines[chapter_start_line:idx]).strip()

            current_clause_num = None
            current_clause_text = []
            chapter_start_line = idx

            chapter_title = _clean_heading(chapter_match.group(1))
            chapter_num = _extract_number(chapter_match.group(1))
            current_chapter = Chapter(number=chapter_num, title=chapter_title)
            chapters.append(current_chapter)
            continue

        if current_chapter is None:
            current_chapter = Chapter(title="前言")
            chapters.append(current_chapter)

        # ── 中文标题识别（无 Markdown 标题的纯文本文档）──────────────────
        if not has_markdown_headings:
            # 章节级：第X部分 / 第X章 / 第X编
            cn_part_match = cn_part_re.match(stripped)
            if cn_part_match:
                _finalize_clause(current_chapter, current_clause_num, current_clause_text)
                if current_chapter is not None and chapters[-1] is current_chapter:
                    current_chapter.text = "\n".join(lines[chapter_start_line:idx]).strip()
                current_clause_num = None
                current_clause_text = []
                chapter_start_line = idx
                full_title = stripped
                current_chapter = Chapter(number=None, title=full_title)
                chapters.append(current_chapter)
                continue

            # 章节级：一、 二、 三、
            cn_num_match = cn_num_re.match(stripped)
            if cn_num_match and len(stripped) < 100:
                _finalize_clause(current_chapter, current_clause_num, current_clause_text)
                if current_chapter is not None and chapters[-1] is current_chapter:
                    current_chapter.text = "\n".join(lines[chapter_start_line:idx]).strip()
                current_clause_num = None
                current_clause_text = []
                chapter_start_line = idx
                current_chapter = Chapter(number=None, title=stripped)
                chapters.append(current_chapter)
                continue

            # 子标题级：（一） （二）
            cn_paren_match = cn_paren_re.match(stripped)
            if cn_paren_match and len(stripped) < 100:
                _finalize_clause(current_chapter, current_clause_num, current_clause_text)
                clause_counter = len(current_chapter.clauses) + 1
                current_chapter.clauses.append(Clause(
                    number=str(clause_counter),
                    text=cn_paren_match.group(1)[:200],
                ))
                continue

            # 子标题级：短行且非正文 → 当作子标题（中文文档中单独成行的短语通常是标题）
            if len(stripped) <= 30 and not stripped.endswith(('。', '；', '：', '.', '!', '?')) \
               and not _clause_match_possible(stripped) and not re.search(r'[，,]', stripped):
                # 启发式：短行 + 不含逗号/句号 → 很可能是子标题，作为 clause 记录
                if current_chapter:
                    _finalize_clause(current_chapter, current_clause_num, current_clause_text)
                    clause_counter = len(current_chapter.clauses) + 1
                    current_chapter.clauses.append(Clause(
                        number=str(clause_counter),
                        text=stripped[:200],
                    ))
                    continue
        h3_match = re.match(r'^###\s+(.+)$', stripped)
        # H2 → 只有当 H2 不是章节级别时才是子标题
        h2_is_sub = chapter_heading_level != 2
        h2_match = re.match(r'^##\s+(.+)$', stripped) if not h3_match and h2_is_sub else None

        if h3_match:
            _finalize_clause(current_chapter, current_clause_num, current_clause_text)
            h3_text = h3_match.group(1)
            num = _extract_number(h3_text)
            if num:
                clause_text = _extract_clause_text_from_heading(h3_text)
                current_chapter.clauses.append(Clause(number=num, text=clause_text[:200]))
                current_clause_num = num
                current_clause_text = []
            else:
                current_clause_num = None
                current_clause_text = [f"【{_clean_heading(h3_text)}】"]
            continue

        if h2_match:
            _finalize_clause(current_chapter, current_clause_num, current_clause_text)
            current_clause_num = None
            current_clause_text = [f"【{_clean_heading(h2_match.group(1))}】"]
            continue

        # 段落级编号匹配
        clause_match = re.match(r'^\s*(\d+(?:\.\d+)*)[）\)\.、]\s+(.*)', stripped)
        if clause_match:
            _finalize_clause(current_chapter, current_clause_num, current_clause_text)
            current_clause_num = clause_match.group(1)
            current_clause_text = [clause_match.group(2).strip()]
            continue

        if not stripped.startswith(("<", "!", "?")):
            clause_num_match = re.match(r'^(\d+(?:\.\d+)+)\s+(.*)', stripped)
            if clause_num_match and len(stripped) < 200:
                _finalize_clause(current_chapter, current_clause_num, current_clause_text)
                current_clause_num = clause_num_match.group(1)
                current_clause_text = [clause_num_match.group(2).strip()]
                continue

        if stripped == "<TABLE>":
            if current_chapter and current_chapter.clauses:
                last = current_chapter.clauses[-1]
                if "[包含表格数据]" not in last.text:
                    last.text = last.text.rstrip() + " [包含表格数据]"
            continue

        if current_clause_num:
            current_clause_text.append(stripped)
        else:
            current_clause_text.append(stripped)

    # 收尾
    _finalize_clause(current_chapter, current_clause_num, current_clause_text)
    if current_chapter is not None:
        current_chapter.text = "\n".join(lines[chapter_start_line:]).strip()

    total = sum(len(ch.clauses) for ch in chapters)

    if not chapters:
        return None

    return DocumentStructure(
        title=chapters[0].title if chapters else "",
        chapters=chapters,
        total_clauses=total,
    )


# ── 内部辅助（保持不变） ─────────────────────────────────────────────────────

def _flatten_html_tables(lines: list[str]) -> list[str]:
    result = []
    in_table = False
    for line in lines:
        has_open = "<table" in line.lower()
        has_close = "</table>" in line.lower()
        if has_open and has_close:
            result.append("<TABLE>")
            continue
        if has_open:
            in_table = True
            result.append("<TABLE>")
            continue
        if has_close:
            in_table = False
            continue
        if not in_table:
            result.append(line)
    return result


def _clause_match_possible(stripped: str) -> bool:
    """快速判断一行是否可能是编号条款（避免被子标题启发式误吞）。"""
    return bool(re.match(r'^\d+(?:\.\d+)*[）\)\.、\s]', stripped))


def _clean_heading(raw: str) -> str:
    text = raw.strip()
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', text).strip()
    return text


def _extract_number(raw: str) -> str:
    text = raw.strip()
    text = re.sub(r'^#{1,3}\s+', '', text)
    m = re.match(r'(\d+(?:\.\d+)*)', text)
    if m:
        return m.group(1)
    return ""


def _extract_clause_text_from_heading(text: str) -> str:
    cleaned = _clean_heading(text)
    return re.sub(r'^\d+(?:\.\d+)*\.?\s*', '', cleaned).strip()


def _finalize_clause(chapter: Optional[Chapter], clause_num: Optional[str], text_lines: list[str]):
    if not clause_num or chapter is None:
        return
    text = " ".join(t for t in text_lines if t).strip()
    if not text:
        return
    if chapter.clauses and chapter.clauses[-1].number == clause_num:
        existing = chapter.clauses[-1].text
        if text not in existing:
            remaining = 200 - len(existing)
            if remaining > 0:
                chapter.clauses[-1].text = existing + " " + text[:remaining]
        return
    chapter.clauses.append(Clause(number=clause_num, text=text[:200]))
