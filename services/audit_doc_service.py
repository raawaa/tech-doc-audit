"""文档解析服务 — 支持 PDF / DOCX / DOC。

解析引擎优先级：
1. MinerU（首选） — 支持文本/扫描件PDF、DOCX，保留表格/布局/层级
2. pdfplumber（PDF降级） — 纯文本提取
3. python-docx（DOCX降级） — 纯文本提取
"""

import os
import re
import json
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from models.audit_document import AuditDocument
import storage.audit_doc_repo as repo


# ── MinerU 探测 ──────────────────────────────────────────────────────────────

_MINERU_BIN: Optional[str] = None

def _find_mineru() -> Optional[str]:
    """查找 MinerU 可执行文件路径（仅通过 PATH 查找，不做任何硬编码）。"""
    import shutil
    try:
        path = shutil.which("mineru")
        if path:
            result = subprocess.run([path, "--version"], capture_output=True, text=True, timeout=5)
            if result.returncode == 0:
                return path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


def _mineru_available() -> bool:
    global _MINERU_BIN
    if _MINERU_BIN is None:
        _MINERU_BIN = _find_mineru()
    return _MINERU_BIN is not None


# ── 文件类型检测 ──────────────────────────────────────────────────────────────

def _detect_file_type(filename: str) -> Optional[str]:
    ext = os.path.splitext(filename)[1].lower()
    mapping = {".pdf": "pdf", ".doc": "doc", ".docx": "docx"}
    return mapping.get(ext)


# ── 文档上传 ──────────────────────────────────────────────────────────────────

def upload_document(original_name: str, content: bytes) -> AuditDocument:
    """上传待审核文档。"""
    file_type = _detect_file_type(original_name)
    if not file_type:
        raise ValueError(f"不支持的文件格式: {original_name}")

    doc = AuditDocument(
        name=original_name,
        original_name=original_name,
        file_type=file_type,
        file_path="",
        status="uploaded",
    )

    doc_dir = repo._doc_dir(doc.id)
    repo._ensure_dir(doc_dir)
    doc.file_path = str(repo._doc_file(doc.id, file_type))
    with open(doc.file_path, "wb") as f:
        f.write(content)

    return repo.save_doc(doc)


def get_document(doc_id: str) -> Optional[AuditDocument]:
    return repo.get_doc(doc_id)


def list_documents() -> list[AuditDocument]:
    return repo.list_docs()


def delete_document(doc_id: str) -> bool:
    return repo.delete_doc(doc_id)


def update_document(doc: AuditDocument) -> AuditDocument:
    return repo.update_doc(doc)


# ── 文档解析 ──────────────────────────────────────────────────────────────────

def parse_document(doc_id: str) -> AuditDocument:
    """解析文档，提取文本和页数。优先使用 MinerU。"""
    doc = repo.get_doc(doc_id)
    if not doc:
        raise ValueError(f"文档不存在: {doc_id}")

    # 已解析过则跳过
    if doc.parsed_content and doc.status == "parsed":
        return doc

    try:
        # 尝试 MinerU
        if _mineru_available() and doc.file_type in ("pdf", "docx"):
            mineru_success = _parse_with_mineru(doc)
            if not mineru_success:
                # MinerU 失败，降级
                _parse_fallback(doc)
        else:
            _parse_fallback(doc)

        doc.status = "parsed"
    except Exception as e:
        doc.status = "failed"
        doc.error_message = str(e)

    return repo.update_doc(doc)


# ── MinerU 解析 ──────────────────────────────────────────────────────────────

def _parse_with_mineru(doc: AuditDocument) -> bool:
    """使用 MinerU 解析文档。返回 True 表示成功，False 表示需要降级。"""
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                [_MINERU_BIN, "-p", doc.file_path, "-o", tmpdir, "-b", "pipeline"],
                capture_output=True, text=True, timeout=300,
            )

            if result.returncode != 0:
                return False

            # 定位输出文件
            base_name = Path(doc.file_path).stem
            md_path = Path(tmpdir) / base_name / "office" / f"{base_name}.md"
            content_list_path = Path(tmpdir) / base_name / "office" / f"{base_name}_content_list.json"

            if not md_path.exists():
                return False

            doc.parsed_content = md_path.read_text(encoding="utf-8")

            # 从 content_list 统计页数
            if content_list_path.exists():
                try:
                    items = json.loads(content_list_path.read_text(encoding="utf-8"))
                    page_nums = set()
                    for item in items:
                        pi = item.get("page_idx")
                        if pi is not None:
                            page_nums.add(int(pi))
                    doc.page_count = max(len(page_nums), 1) if page_nums else None
                except Exception:
                    pass

            if doc.page_count is None:
                doc.page_count = _estimate_page_count(doc.parsed_content)

            return True

    except Exception:
        return False


def _estimate_page_count(content: str) -> int:
    """从 Markdown 内容估算页数。"""
    # 计算 HTML 表格数量（通常一页 1-2 个表格）
    table_count = content.count("<table>")
    if table_count > 0:
        return max(1, table_count)
    # 按 3000 字一页估算
    char_count = len(content)
    return max(1, char_count // 3000)


# ── 降级解析 ──────────────────────────────────────────────────────────────────

def _parse_fallback(doc: AuditDocument):
    """降级方案：用 pdfplumber / python-docx。"""
    if doc.file_type == "pdf":
        _parse_pdf_fallback(doc)
    elif doc.file_type == "docx":
        _parse_docx_fallback(doc)
    else:
        doc.parsed_content = f"[暂不支持解析 .{doc.file_type} 文件]"
        doc.page_count = None


def _parse_pdf_fallback(doc: AuditDocument):
    import pdfplumber
    text_parts = []
    with pdfplumber.open(doc.file_path) as pdf:
        doc.page_count = len(pdf.pages)
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    doc.parsed_content = "\n\n".join(text_parts)


def _parse_docx_fallback(doc: AuditDocument):
    import docx
    text_parts = []
    try:
        document = docx.Document(doc.file_path)
        for para in document.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        for table in document.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    text_parts.append(" | ".join(row_texts))
        doc.parsed_content = "\n".join(text_parts)
        doc.page_count = max(1, len(text_parts) // 30)
    except Exception as e:
        doc.parsed_content = f"[Word 文档解析失败: {e}]"
        doc.page_count = None
