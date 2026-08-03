"""PDF / DOCX 文档解析的唯一入口（PRD #29 / V2）。

``parse_document(file_path) -> ParseResult`` 是 KB 文档导入流水线的起点。
一次解析返回 {by_page, full_text, layout}，所有下游（按页文本存储 / 向量索引 / 文本搜索 / PDF 跳转）
从同一份数据消费——避免历史上双解析器（``extract_text`` + ``extract_text_by_page``）导致的不一致。

降级链（P1 数据层 #32 V1 已落 cache）：
  1. PDF: PaddleOCR-VL-1.6（带缓存，命中即跳过 OCR 配额）
  2. PDF 缓存未命中 → PaddleOCR 重新推理 → 落缓存
  3. PaddleOCR 不可用/失败 → 提取页面 markdown 聚合到 full_text + by_page=单页
  4. PaddleOCR/MinerU 全失败 → pdfplumber 流式逐页抽取
  5. 非 PDF（DOCX / MD）→ 单页 full_text，无 OCR

文本层 PDF 另解(issue #99/#105):`_pymupdf_parse(file_path) -> ParseResult`
按 PyMuPDF text-block + image-block + tables(block_label="table")产
ParseResult,**不**接管路由(routing 仍由 `_parse_pdf` 负责)。
"""
from __future__ import annotations

import json
import os
import re
import time
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Optional

from core.logger import get_logger

_logger = get_logger(__name__)


# ── 数据类 ──────────────────────────────────────────────────────────────────────


@dataclass
class PageText:
    """按页纯文本。page 编号 0-based。"""
    page: int
    text: str


@dataclass
class Block:
    """PaddleOCR ``prunedResult.parsing_res_list`` 一项的归一化坐标视图。

    bbox / polygon 以归一化坐标 ``[x1/W, y1/H, x2/W, y2/H]`` 存储（0-1 浮点），
    与 PDF 渲染分辨率解耦；page 级别同时存原始 width / height 便于换算。
    """
    block_label: str = ""
    block_content: str = ""
    bbox_norm: list[float] = field(default_factory=list)
    polygon_norm: list[list[float]] = field(default_factory=list)
    block_order: int = 0


@dataclass
class PageLayout:
    """单个 PDF 页面的版面信息（含归一化坐标 blocks）。"""
    page: int
    width: int = 0
    height: int = 0
    blocks: list[Block] = field(default_factory=list)


@dataclass
class ParseResult:
    """一次解析产出的结构化结果。"""
    by_page: list[PageText]
    full_text: str
    layout: list[PageLayout]

    def to_dict(self) -> dict:
        """序列化供 cache / pages_store 落盘。"""
        return {
            "by_page": [asdict(p) for p in self.by_page],
            "full_text": self.full_text,
            "layout": [
                {
                    "page": pl.page,
                    "width": pl.width,
                    "height": pl.height,
                    "blocks": [asdict(b) for b in pl.blocks],
                }
                for pl in self.layout
            ],
        }

    @classmethod
    def from_dict(cls, data: dict) -> "ParseResult":
        """从 cache / pages_store 反序列化。"""
        by_page = [PageText(**p) for p in data.get("by_page", [])]
        layout = [
            PageLayout(
                page=pl.get("page", 0),
                width=pl.get("width", 0),
                height=pl.get("height", 0),
                blocks=[Block(**b) for b in pl.get("blocks", [])],
            )
            for pl in data.get("layout", [])
        ]
        return cls(
            by_page=by_page,
            full_text=data.get("full_text", ""),
            layout=layout,
        )


# ── 解析入口 ────────────────────────────────────────────────────────────────────


_PADDLEOCR_API_URL = os.environ.get("PADDLEOCR_API_URL", "").rstrip("/")
_PADDLEOCR_API_TOKEN = os.environ.get("PADDLEOCR_API_TOKEN", "").strip()
_PADDLEOCR_MODEL = os.environ.get("PADDLEOCR_MODEL", "PaddleOCR-VL-1.6")


def _paddleocr_available() -> bool:
    return bool(_PADDLEOCR_API_TOKEN and _PADDLEOCR_API_URL)


def _pymupdf_available() -> bool:
    try:
        import pymupdf  # noqa: F401
    except Exception:
        return False
    return True


def _is_text_layer_pdf(file_path: str) -> bool:
    """Return whether a PDF contains at least 20 characters of text."""
    if not file_path or Path(file_path).suffix.lower() != ".pdf":
        return False

    try:
        import pymupdf

        with pymupdf.open(file_path) as document:
            text = "".join(page.get_text() for page in document)
        return len(text.strip()) >= 20
    except Exception as e:
        _logger.debug("text-layer detection failed for %s: %s", file_path, e)
        return False


def parse_document(file_path: str, *, use_cache: bool = True) -> ParseResult:
    """解析单份文档，返回 ParseResult。

    PDF: 缓存命中 → 直接反序列化；未命中 → PaddleOCR → 缓存。
    非 PDF: 单页 ParseResult。
    失败: 返回 by_page=[PageText(0, full_text)] 的兜底 ParseResult，full_text 可能为空。

    Args:
        file_path: 文档绝对路径。
        use_cache: True（默认）→ 查/写 ``core.paddleocr_cache``；False → 强制重新解析。
    """
    if not file_path or not Path(file_path).exists():
        return _empty_result()

    ext = Path(file_path).suffix.lower()

    if ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    if ext in (".md", ".markdown", ".txt"):
        return _parse_plain_text(file_path)
    if ext != ".pdf":
        # 其他格式：视作文本读取
        return _parse_plain_text(file_path)

    return _parse_pdf(file_path, use_cache=use_cache)


# ── PDF 路径（含缓存）────────────────────────────────────────────────────────────


def _parse_pdf(file_path: str, *, use_cache: bool) -> ParseResult:
    cached: Optional[dict] = None
    if use_cache:
        from core.paddleocr_cache import get_cached
        cached = get_cached(file_path)

    if cached is not None:
        return ParseResult.from_dict(cached)

    result, source = _paddleocr_parse(file_path)

    if use_cache and result.by_page:
        try:
            from core.paddleocr_cache import save_cached
            save_cached(file_path, result.to_dict(), source=source)
        except Exception as e:
            _logger.warning("parse_document: cache save failed for %s: %s", file_path, e)

    return result


def _paddleocr_parse(file_path: str) -> tuple[ParseResult, str]:
    """调 PaddleOCR API → 落 ParseResult；不可用 / 失败 → 降级。

    返回 ``(result, source)``: source 是 cache 元数据，标识结果来自哪个解析器
    （见 ``paddleocr_cache.save_cached`` 的 source 参数说明，issue #57）。
    """
    if not _paddleocr_available():
        result = _pdf_fallback(file_path)
        return result, "fallback_pdfplumber"

    try:
        result = _paddleocr_call(file_path)
        if not result.full_text or len(result.full_text) < 20:
            _logger.info("paddleocr returned empty for %s, retrying with orientation classify", file_path)
            result = _paddleocr_call(file_path, orientation_classify=True)
        return result, "paddleocr"
    except Exception as e:
        _logger.warning("paddleocr_parse failed for %s: %s", file_path, e)
        return _pdf_fallback(file_path), "fallback_pdfplumber"


def _paddleocr_call(file_path: str, orientation_classify: bool = False) -> ParseResult:
    """实际的 PaddleOCR-VL-1.6 调用流程。失败时抛 Exception。"""
    import requests  # 延迟 import，允许离线测试

    headers = {"Authorization": f"bearer {_PADDLEOCR_API_TOKEN}"}

    # 提交 job
    data = {
        "model": _PADDLEOCR_MODEL,
        "optionalPayload": json.dumps({
            "useDocOrientationClassify": orientation_classify,
            "useDocUnwarping": False,
            "useChartRecognition": False,
        }),
    }
    with open(file_path, "rb") as f:
        resp = requests.post(
            _PADDLEOCR_API_URL, headers=headers, data=data,
            files={"file": f}, timeout=120,
        )
    resp.raise_for_status()
    job_id = resp.json()["data"]["jobId"]

    # 轮询
    deadline = time.monotonic() + 600
    jsonl_url = ""
    while time.monotonic() < deadline:
        try:
            r = requests.get(f"{_PADDLEOCR_API_URL}/{job_id}", headers=headers, timeout=30)
            r.raise_for_status()
            j = r.json()["data"]
            state = j["state"]
            if state == "done":
                jsonl_url = j["resultUrl"]["jsonUrl"]
                break
            if state == "failed":
                raise RuntimeError(f"paddleocr job failed: {j.get('errorMsg', '?')}")
        except RuntimeError:
            raise
        except Exception:
            pass
        time.sleep(5)

    if not jsonl_url:
        raise RuntimeError("paddleocr timeout")

    # 取 JSONL
    r = requests.get(jsonl_url, timeout=120)
    r.raise_for_status()

    return _paddleocr_jsonl_to_parse_result(r.text)


def _paddleocr_jsonl_to_parse_result(jsonl_text: str) -> ParseResult:
    """解析 PaddleOCR JSONL 响应 → ParseResult。

    每行对应一页的 layoutParsingResults；从 ``markdown.text`` 取页面文本，
    从 ``prunedResult.parsing_res_list`` 取归一化 block 坐标。
    标题层级修复（HeadingProcessor）应用到 full_text。

    V7.1 None-safety：宽高兜底 ``res.width/height`` → ``prunedResult.width/height``
    → ``prunedResult.image_size`` → 0；任一维度缺省时产出的 block 不做坐标归一化
    （坐标为空 list），但 block 本身仍保留（label / content / block_order 不丢）。
    多页打包：单 JSONL 行 ``layoutParsingResults`` 按出现顺序累加 page_order。
    """
    page_texts: list[PageText] = []
    page_layouts: list[PageLayout] = []
    page_order = 0

    for line in jsonl_text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        try:
            payload = json.loads(line)
        except json.JSONDecodeError:
            continue
        result = payload.get("result") or {}
        for res in result.get("layoutParsingResults") or []:
            md_text = (res.get("markdown", {}) or {}).get("text", "") or ""
            page_texts.append(PageText(page=page_order, text=md_text.strip()))

            pruned = res.get("prunedResult") or {}
            blocks = _extract_blocks(pruned)

            # 宽高：res → prunedResult → image_size → 0
            width_raw = res.get("width")
            if width_raw is None:
                width_raw = pruned.get("width")
            height_raw = res.get("height")
            if height_raw is None:
                height_raw = pruned.get("height")
            width = int(width_raw or 0)
            height = int(height_raw or 0)
            if not width or not height:
                img_w, img_h = _page_dims(pruned.get("image_size"))
                width = width or img_w
                height = height or img_h

            page_layouts.append(PageLayout(
                page=page_order,
                width=width,
                height=height,
                blocks=blocks,
            ))
            page_order += 1

    full_md = "\n\n".join(p.text for p in page_texts if p.text)
    if full_md:
        full_md = _normalize_headings(full_md)

    return ParseResult(by_page=page_texts, full_text=full_md, layout=page_layouts)


def _extract_blocks(pruned: dict) -> list[Block]:
    """``prunedResult.parsing_res_list`` → ``[Block]``，bbox / polygon 归一化到 0-1。

    V7.1 None-safety：PaddleOCR 真实响应里 ``block_order`` / ``block_bbox`` /
    ``block_polygon`` 偶尔为 ``None``；宽高可能在 ``prunedResult.width`` /
    ``prunedResult.image_size`` 两层都缺。任意字段为 None 时降级到空坐标（block
    仍产出，命中坐标丢失），不能崩。

    宽高优先级：``prunedResult.width/height`` → ``prunedResult.image_size`` → 0。
    （``res.width/height`` 已在 _paddleocr_jsonl_to_parse_result 合并到 prunedResult 层）
    """
    raw_blocks = pruned.get("parsing_res_list") or []
    W = int(pruned.get("width") or 0)
    H = int(pruned.get("height") or 0)
    if not W or not H:
        img_w, img_h = _page_dims(pruned.get("image_size"))
        W = W or img_w
        H = H or img_h

    blocks: list[Block] = []
    for i, b in enumerate(raw_blocks):
        # 容忍 None：缺字段时退到空字符串 / 空 list；block 仍保留（命中坐标系丢失）
        bbox = _coerce_bbox(b.get("block_bbox") or b.get("bbox"))
        polygon_raw = (
            b.get("block_polygon")
            or b.get("polygon")
            or b.get("block_polygon_points")
            or []
        )
        if polygon_raw is None:
            polygon_raw = []

        if W and H:
            norm_bbox = (
                [bbox[0] / W, bbox[1] / H, bbox[2] / W, bbox[3] / H]
                if len(bbox) == 4 else []
            )
            norm_polygon = [
                [(float(p[0]) / W), (float(p[1]) / H)]
                for p in polygon_raw if isinstance(p, (list, tuple)) and len(p) >= 2
            ]
        else:
            norm_bbox, norm_polygon = [], []

        # block_order 缺省 / None / 非法时退回索引 i
        order_raw = b.get("block_order", i)
        try:
            block_order = int(order_raw) if order_raw is not None else i
        except (TypeError, ValueError):
            block_order = i

        blocks.append(Block(
            block_label=str(b.get("block_label", "") or b.get("label", "") or ""),
            block_content=str(b.get("block_content", "") or b.get("content", "") or ""),
            bbox_norm=norm_bbox,
            polygon_norm=norm_polygon,
            block_order=block_order,
        ))
    return blocks


def _page_dims(image_size) -> tuple[int, int]:
    """``prunedResult.image_size`` → (W, H)。"""
    if isinstance(image_size, (list, tuple)) and len(image_size) >= 2:
        try:
            return int(image_size[0]), int(image_size[1])
        except (TypeError, ValueError):
            pass
    return 0, 0


def _coerce_bbox(raw) -> list:
    """Accept [x1,y1,x2,y2] / (x1,y1,x2,y2) / None → flat list[float]."""
    if raw is None:
        return []
    if isinstance(raw, (list, tuple)):
        try:
            return [float(v) for v in raw[:4]]
        except (TypeError, ValueError):
            return []
    return []


def _normalize_headings(text: str) -> str:
    """标题层级修复（保留 ``core.heading_processor.HeadingProcessor`` 失败兜底）。"""
    try:
        from core.heading_processor import HeadingProcessor
        return HeadingProcessor().rebuild_from_md(text)
    except Exception:
        return text


# ── PyMuPDF 文本层解析（issue #99/#105）──────────────────────────────────────────


def _pymupdf_parse(file_path: str) -> ParseResult:
    """PyMuPDF 文本层 PDF 解析(issue #99 主路径之一,PRD 来自 #105)。

    关键陷阱(**PyMuPDF 文档示例不会告诉你**):``page.find_tables().tables[i].bbox``
    是 4-tuple ``(x0, y0, x1, y1)``——**不是** ``fitz.Rect``。原 subagent 代码曾用
    ``tb.x0`` 被 silent ``except: pass`` 吞掉 → emit 0 table blocks。
    本实现直接读取 4-tuple。

    异常:损坏 / 加密 PDF 时 PyMuPDF 自身抛 ``pymupdf.FileDataError`` /
    ``pymupdf.FilePasswordError``——**不在此函数内 swallow**,让路由层决定
    fallback / 报错(pymupdf 自身已经抛得足够清晰)。

    坐标系:bbox / polygon 以归一化坐标 ``[x/W, y/H, x/W, y/H]`` 存储(0-1)。
    """
    try:
        import pymupdf
    except ImportError:
        # 没有 wheel → 抛清晰错误,让上层看到
        raise RuntimeError(
            "pymupdf wheel not installed; install pymupdf==1.28.0 to use "
            "_pymupdf_parse (see issue #99 / #105)"
        )

    by_page: list[PageText] = []
    layouts: list[PageLayout] = []
    full_text_parts: list[str] = []
    block_order = 0  # 跨页单调递增

    with pymupdf.open(file_path) as doc:
        for page_index, page in enumerate(doc):
            W = float(page.rect.width)
            H = float(page.rect.height)
            page_text = page.get_text("text") or ""
            by_page.append(PageText(page=page_index, text=page_text))
            if page_text:
                full_text_parts.append(page_text)

            blocks = _pymupdf_page_blocks(page, W, H, block_order)
            layouts.append(PageLayout(
                page=page_index,
                width=int(W),
                height=int(H),
                blocks=blocks,
            ))
            block_order += len(blocks)

    full_text = "\n\n".join(full_text_parts)
    if full_text:
        full_text = _normalize_headings(full_text)
    return ParseResult(by_page=by_page, full_text=full_text, layout=layouts)


def _pymupdf_page_blocks(
    page, W: float, H: float, start_order: int
) -> list[Block]:
    """单页 PyMuPDF → ``Block[]``(text + image + table,单页内 block_order 从
    ``start_order`` 起递增)。

    顺序:text(``get_text("dict")`` 中的 type=0)+ image(同 type=1),按 PDF 源顺序;
    之后 tables(``find_tables().tables``)按检测顺序追加在文本/图像 blocks 之后。
    跨多页时,block_order 连续递增(上游 ``_pymupdf_parse`` 在 page 间累加)。
    """
    blocks: list[Block] = []
    if not (W > 0 and H > 0):
        return blocks

    # ── text + image:``page.get_text("dict")`` 给出 PDF 源顺序的 layout ─────────
    # 注意:**不**在这里 swallow 异常——issue #105 明确要求"Pymupdf 抛异常不在
    # 函数内 swallow,让调用方决定"。只有上游 ``pymupdf.open()`` 失败不在此 catch。
    page_dict = page.get_text("dict")

    block_order = start_order
    for raw in page_dict.get("blocks", []) or []:
        btype = raw.get("type", 0)
        bbox = _coerce_bbox(raw.get("bbox"))
        if len(bbox) != 4:
            continue
        norm_bbox, polygon_norm = _norm_bbox_polygon(bbox, W, H)

        if btype == 1:
            # image block:内容字段写 "[image]" 标记(与 PaddleOCR image block 兼容)
            blocks.append(Block(
                block_label="image",
                block_content="[image]",
                bbox_norm=norm_bbox,
                polygon_norm=polygon_norm,
                block_order=block_order,
            ))
        else:
            # text block:拼接所有 spans 的 text
            text_chunks: list[str] = []
            for line in raw.get("lines", []) or []:
                for span in line.get("spans", []) or []:
                    t = span.get("text")
                    if t:
                        text_chunks.append(t)
            blocks.append(Block(
                block_label="text",
                block_content="".join(text_chunks),
                bbox_norm=norm_bbox,
                polygon_norm=polygon_norm,
                block_order=block_order,
            ))
        block_order += 1

    # ── tables:``find_tables().tables`` → 4-tuple bbox,**不是 fitz.Rect** ───────
    # 同样不 swallow:PyMuPDF 抛任何异常都向上传,路由层决定 fallback。
    for tbl in (page.find_tables().tables or []):
        # ``tbl.bbox`` 是 ``(x0, y0, x1, y1)``;原 1-line bug 就是想抓
        # ``.x0`` 被 ``AttributeError`` swallow。这里直接读 4-tuple。
        table_bbox = _coerce_bbox(tbl.bbox)
        if len(table_bbox) != 4:
            continue
        norm_bbox, polygon_norm = _norm_bbox_polygon(table_bbox, W, H)
        # 把 cells 文本拼成 block_content(给下游 consumers 一个概览)
        cell_parts: list[str] = []
        for row in (tbl.cells or []):
            row_vals: list[str] = []
            if isinstance(row, (list, tuple)):
                row_vals = [str(c) if c is not None else "" for c in row]
            cell_parts.append(" | ".join(row_vals))
        blocks.append(Block(
            block_label="table",
            block_content="\n".join(cell_parts),
            bbox_norm=norm_bbox,
            polygon_norm=polygon_norm,
            block_order=block_order,
        ))
        block_order += 1

    return blocks


def _norm_bbox_polygon(
    bbox, W: float, H: float
) -> tuple[list[float], list[list[float]]]:
    """bbox ``(x0, y0, x1, y1)`` → ``(bbox_norm 4-list, polygon_norm 4-corner)``。
    坐标全部 clamp 到 ``[0,1]``(可能 page.cropbox > mediabox 时出现轻微越界)。
    """
    x0 = float(bbox[0]) / W
    y0 = float(bbox[1]) / H
    x1 = float(bbox[2]) / W
    y1 = float(bbox[3]) / H
    norm_bbox = [max(0.0, min(1.0, v)) for v in (x0, y0, x1, y1)]
    polygon = [
        [norm_bbox[0], norm_bbox[1]],
        [norm_bbox[2], norm_bbox[1]],
        [norm_bbox[2], norm_bbox[3]],
        [norm_bbox[0], norm_bbox[3]],
    ]
    return norm_bbox, polygon


# ── PDF 降级:pdfplumber 流式逐页 ───────────────────────────────────────────────


def _pdf_fallback(file_path: str) -> ParseResult:
    """PaddleOCR 不可用 / 失败：走 pdfplumber 流式逐页抽取，落 ParseResult。"""
    try:
        import pdfplumber
    except ImportError:
        return _empty_result()

    page_texts: list[PageText] = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text() or ""
                page_texts.append(PageText(page=i, text=t))
                page.flush_cache()
    except Exception as e:
        _logger.warning("pdfplumber fallback failed for %s: %s", file_path, e)
        return _empty_result()

    full_text = "\n\n".join(p.text for p in page_texts if p.text)
    full_text = _normalize_headings(full_text)
    return ParseResult(by_page=page_texts, full_text=full_text, layout=[])


# ── 非 PDF 路径 ────────────────────────────────────────────────────────────────


def _parse_docx(file_path: str) -> ParseResult:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return _empty_result()
    try:
        parts = [p.text for p in DocxDocument(file_path).paragraphs if p.text and p.text.strip()]
    except Exception as e:
        _logger.warning("docx parse failed for %s: %s", file_path, e)
        return _empty_result()
    text = "\n".join(parts)
    return ParseResult(
        by_page=[PageText(page=0, text=text)] if text else [],
        full_text=text,
        layout=[],
    )


def _parse_plain_text(file_path: str) -> ParseResult:
    try:
        text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        _logger.warning("plain text read failed for %s: %s", file_path, e)
        return _empty_result()
    return ParseResult(
        by_page=[PageText(page=0, text=text)] if text else [],
        full_text=text,
        layout=[],
    )


def _empty_result() -> ParseResult:
    """PaddleOCR + pdfplumber 都失败时的兜底：空结构。"""
    return ParseResult(by_page=[], full_text="", layout=[])
